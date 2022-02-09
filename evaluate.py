#!/usr/bin/python3

from collections import OrderedDict
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from glob import glob
from joblib import dump, load
from matplotlib import cm
from matplotlib import pyplot
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
from matplotlib.transforms import Bbox
from PIL import Image
from PIL import ImageDraw, ImageFont
from scipy.stats import spearmanr, pearsonr
from typing import Dict, Any
import copy
import cv2
import hashlib
import itertools
import json
import math
import matplotlib
import matplotlib.colors as colors
import matplotlib.pyplot as plt
import multiprocessing
import numpy as np
import os
import pandas as pd
import pathlib
import pickle
import pylab
import random
import scipy.cluster.hierarchy as sch
import seaborn as sns
import shutil
import sys
import tempfile
import time

# delong
import rpy2
import rpy2.robjects as robjects
from rpy2.robjects.packages import importr
from rpy2.robjects import pandas2ri
pandas2ri.activate()


from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *
from contextlib import contextmanager





### parameters
cFigNumber = 1
document = None

def getResults (dList):
    mlflow.set_tracking_uri(TrackingPath)
    if os.path.exists("./results/results.feather") == False:
        results = []
        for d in dList:
            current_experiment = dict(mlflow.get_experiment_by_name(d))
            experiment_id = current_experiment['experiment_id']
            runs = MlflowClient().search_runs(experiment_ids=experiment_id, max_results=50000)
            for r in runs:
                row = r.data.metrics
                row["UUID"] = r.info.run_uuid
                row["Model"] = r.data.tags["Version"]
                row["Parameter"] = r.data.tags["pID"]

                # stupid naming error
                row["Parameter"] = row["Parameter"]
                row["Model"] = row["Model"]

                row["FSel"], row["Clf"] = row["Model"].split("_")
                row["Dataset"] = d

                row["nFeatures"] = eval(row["Parameter"])[row["FSel"]]["nFeatures"]

                row["Path"] = os.path.join(TrackingPath,  str(experiment_id), str(r.info.run_uuid), "artifacts")
                results.append(row)

                # read timings
                apath = os.path.join(row["Path"], "timings.json")
                with open(apath) as f:
                    expData = json.load(f)
                row.update(expData)

                # read AUCs
                apath = os.path.join(row["Path"], "aucStats.json")
                with open(apath) as f:
                    aucData = json.load(f)
                row.update(aucData)

        results = pd.DataFrame(results)
        print ("Pickling results")
        pickle.dump (results, open("./results/results.feather","wb"))
    else:
        print ("Restoring results")
        results = pickle.load(open("./results/results.feather", "rb"))

    return results




def delongTest (predsX, predsY):
    pROC = importr('pROC')

    Y = predsX["y_true"].values
    scoresA = predsX["y_pred"].values
    scoresB = predsY["y_pred"].values
    rocA = pROC.roc (Y, scoresA)
    rocB = pROC.roc (Y, scoresB)

    aucA = pROC.auc(Y, scoresA)
    aucB = pROC.auc(Y, scoresB)
    #print ("AUC A:" + str(aucA))
    #print ("AUC B:" + str(aucB))
    robjects.globalenv['rocA'] = rocA
    robjects.globalenv['rocB'] = rocB

    z = rpy2.robjects.packages.reval ("library(pROC);z = roc.test(rocA, rocB, method= 'delong', progress='none'); p = z$p.value")
    z = robjects.r.z
    p = robjects.r.p[0]
    return p



def getDeLongTest (dList, results):
    # check cache..
    if os.path.exists("./results/delongTests.joblib") == False:
        rMatList = {}
        for d in dList:
            print(d)
            fSels = sorted(list(set(results["FSel"].values)))
            nList = sorted(list(set(results["nFeatures"].values)))
            clfs = sorted(list(set(results["Clf"].values)))
            rMat = np.zeros( (len(clfs), len(fSels) ) )
            rMat = pd.DataFrame(rMat, index = clfs, columns = fSels)

            # get best overall
            aTable = results.query("Dataset == @d and FSel in @fSels")
            aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
            best = aTable.iloc[0].copy()

            # before we start, test if best model is different from random?
            predsX = pd.read_csv(os.path.join(best["Path"], "preds.csv"))
            predsY = predsX.copy()
            predsY["y_pred"] = 0.5
            print ("Testing if best model is better than random:")
            p = delongTest (predsX, predsY)

            if p < 0.05:
                print ("Yes, p = ", p)
            else:
                print ("No, p = ", p)

            for c in clfs:
                for f in fSels:
                    aTable = results.query("Dataset == @d and FSel == @f and Clf == @c")
                    aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
                    cur = aTable.iloc[0].copy()

                    # load both preds
                    predsX = pd.read_csv(os.path.join(best["Path"], "preds.csv"))
                    predsY = pd.read_csv(os.path.join(cur["Path"], "preds.csv"))

                    p = delongTest (predsX, predsY)
                    if p < 0.05:
                        pass
                    else:
                        rMat.at[c,f] = p
            rMatList[d] = rMat
        dump(rMatList, "./results/delongTests.joblib")
    else:
        print ("Restoring delong results")
        rMatList = load( "./results/delongTests.joblib")
    return rMatList



def plot_DataHisto (dList):
    DPI  = 300
    fig, ax = plt.subplots(4,4, figsize = (25, 20), dpi = DPI)
    N = len(dList)
    palette = sns.color_palette("hls", N+N//3)[:N]
    for fidx, d in enumerate(dList):
        X, y = datasets[d]

        M = X.corr().values
        mask = np.triu(M*0+1, k =1)
        v = np.extract(mask, M)

        fidx_y= fidx % 4
        fidx_x = fidx//4
        doHistoPlot (v, d, "./results/Data_Hist_" + d + ".png", ax = ax[fidx_x][fidx_y], color = palette[fidx], fig = fig)

    # use last one, does not matter which one actually
    data = eval (d+"().getData('./data/')")
    y = data["Target"]
    X = data.drop(["Target"], axis = 1)
    X, y = preprocessData (X, y)

    arrays = [np.random.normal(loc=0, scale=1, size=(X.shape[0])) for s in range(X.shape[1])]
    X = np.vstack(arrays).T
    X = pd.DataFrame(X)
    M = X.corr().values
    mask = np.triu(M*0+1, k =1)
    f = np.extract(mask, M)
    doHistoPlot (f, "Normal", "./results/Data_Hist_Normal.png", ax = ax[3][3], color = "black", range = True, fig = fig)


    # remove unused plot
    for x, y in [ (3,2)]:
        ax[x][y] .spines['right'].set_visible(False)
        ax[x][y] .spines['top'].set_visible(False)
        ax[x][y] .spines['bottom'].set_visible(False)
        ax[x][y] .spines['left'].set_visible(False)
        ax[x][y] .spines['left'].set_visible(False)
        ax[x][y].xaxis.set_visible(False)
        ax[x][y].yaxis.set_visible(False)

    plt.tight_layout(pad=3.0)
    fig.savefig("./paper/Figure_2.png", facecolor = 'w', bbox_inches='tight')
    plt.close('all')
    plt.rc('text', usetex=False)
    pass




def countModels (dList, results):
    rMatList = getDeLongTest (dList, results)
    cnts = {}
    for d in rMatList:
        z = rMatList[d] >= 0.05
        cnts[d] = {"Count": int(np.sum(z.values)) - 1, # because best model does not count
                    "AUC": np.round (np.max(results.query(' Dataset == @d')["AUC"]), 2) }

    df = pd.DataFrame(cnts).T[["AUC", "Count"]]
    df["Count"] = df["Count"].astype(np.uint32)
    df = df.sort_values(["AUC"], ascending = False)
    df.to_csv("./paper/Table_4.csv")


    # dont care which f,
    f = "MIM"
    z = results.query(' Dataset == @d and FSel == @f')
    nDataSets= len(set(results["Dataset"]))
    print ("#Datasets", nDataSets)
    print ("#Models per Dataset", results.shape[0]/nDataSets)
    nFSel= len(set(results["FSel"]))
    print ("#Models per FSel and Dataset", results.shape[0]/nDataSets/nFSel)
    nClf = len(set(results["Clf"]))
    print ("#Classifier", nClf)

    # number of features=1,2,..64= 7
    nF = len(set(z["nFeatures"]))
    print ("Have", nF, "number of features")

    # number of hyperparameters
    cf = z["Clf"].value_counts()/nF
    print ("Have for each FSel", sum(cf.values), "models/hyperparameters")

    total = rMatList[d].shape[0]*    rMatList[d].shape[1] - 1  # best model doesnt really count
    print ("Total stat.eq. models over all datasets:", np.mean(df["Count"]), "/", total)
    print ("Percentage of stat eq. models per dataset:", np.mean(df["Count"])/total)
    total = total*len(rMatList)
    print ("Total stat.eq. models over all datasets:", np.sum(df["Count"]), "/", total)
    print ("Percentage of stat.eq. models over all datasets:", np.sum(df["Count"])/total)


def plot_modelAUCs (dList, results):
    print ("Plotting model AUCs")
    rMatList = getDeLongTest (dList, results)
    sTable = []
    for d in rMatList.keys():
        # no idea what pandas fct to use..
        for f in rMatList[d].keys():
            for c in rMatList[d].index:
                sTable.append({"Dataset": d, "FSel": f, "Clf": c, "p": rMatList[d].at[c,f]})
    sTable = pd.DataFrame(sTable)
    for (i, (idx, row)) in enumerate(sTable.iterrows()):
        d, f, c = row["Dataset"], row["FSel"], row["Clf"]
        aTable = results.query ("Dataset == @d and FSel == @f and Clf == @c")
        aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
        cur = aTable.iloc[0].copy()
        sTable.at[idx, "AUC"] = cur["AUC"]
    sTable["Statistically Similar"] = sTable["p"] >= 0.05

    if 1 == 1:
            DPI  = 300
            fig, ax = plt.subplots(figsize = (15, 10), dpi = DPI)
            sns.set(style='white')
            #strange but well
            # nFSel = len(set([k[0] for k in z.index]))
            palette = sns.color_palette("hls", 8)[0::4]
            sns.stripplot(x="AUC", y="Dataset", jitter = 0.25, data=sTable, palette = palette, hue = "Statistically Similar")

            plt.xticks(fontsize=20)
            plt.yticks(fontsize=20)
            plt.ylabel('Dataset', fontsize = 22, labelpad = 12)
            plt.xlabel('AUC-ROC', fontsize= 22, labelpad = 12)
            plt.setp(ax.get_legend().get_texts(), fontsize='16') # for legend text
            plt.setp(ax.get_legend().get_title(), fontsize='20') # for legend title
            #ax.set_xticks(nList[1:])#, rotation = 0, ha = "right", fontsize = 22)

            ax.xaxis.tick_bottom()
            ax.yaxis.tick_left()
            plt.tight_layout()
            fig.savefig("./paper/Figure_1.png", facecolor = 'w', bbox_inches='tight')
    print ("Done: Plotting model AUCs")
    pass



def stability (u, v):
    assert (len(u) == len(v))
    m = len(u)
    SC = 0
    for i in range(m):
        for j in range(i+1,m):
            coef, p = pearsonr(u[i,:], v[j,:])
            SC = SC + coef
    SC = 2/(m*(m-1))*SC
    return SC


def getFPattern (model):
    z = []
    for m in range(nCV):
        apath = os.path.join(model["Path"], "FPattern_" + str(m) + ".json")
        with open(apath) as f:
            expData = json.load(f)
        z.append(list(expData.values()))
    z = np.asarray(z)
    return z


def getStability (model):
    z = getFPattern (model)
    SC = stability (z, z)
    return SC


def plot_Stability_Curves (dList, results):
    global document;
    global cFigNumber;

    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 2')
    document.add_paragraph(' ')
    document.add_heading('Stability vs. number of features', level = 2)
    document.add_paragraph(' ')

    # check cache..
    if os.path.exists("./results/stability.joblib") == False:
        fTable = []
        for d in dList:
            fSels = sorted(list(set(results["FSel"].values)))
            nList = sorted(list(set(results["nFeatures"].values)))
            for f in fSels:
                for n in nList:
                    aTable = results.query("Dataset == @d and FSel == @f and nFeatures == @n").copy()
                    for (i, (idx, row)) in enumerate(aTable.iterrows()):
                        stab = getStability(row)
                        aTable.at[idx, "Stability"] = stab
                    # should be the same, but if there was any kind of numerical error,
                    # taking median makes more sense
                    fTable.append({"Dataset": d, "Feature Selection": f, "N": n, "Stability": np.median(aTable["Stability"]) })

        dump(fTable, "./results/stability.joblib")
    else:
        print ("Restoring stability results")
        fTable = load( "./results/stability.joblib")

    # group by dataset and take mean
    df = pd.DataFrame(fTable)
    z = df.groupby(["Feature Selection", "N"]).median(["Stability"])
    nList = sorted(list(set(df["N"].values)))

    def doPlot (z, fname):
            DPI  = 300
            fig, ax = plt.subplots(figsize = (10, 10), dpi = DPI)
            sns.set(style='white')
            #strange but well
            nFSel = len(set([k[0] for k in z.index]))
            palette = sns.color_palette("hls", nFSel+1)[0:nFSel]
            palette[2] = (0.9, 0.9, 0.2)
            line=sns.lineplot(x="N", y="Stability",hue="Feature Selection", palette = palette, marker="o", data=z, linewidth = 4)
            plt.xticks(fontsize=20)
            plt.yticks(fontsize=20)
            plt.ylabel('Mean Stability', fontsize = 22, labelpad = 12)
            plt.xlabel('Number of Selected Features', fontsize= 22, labelpad = 12)
            plt.setp(ax.get_legend().get_texts(), fontsize='16') # for legend text
            plt.setp(ax.get_legend().get_title(), fontsize='20') # for legend title
            ax.set_xticks(nList[1:])#, rotation = 0, ha = "right", fontsize = 22)

            # set the linewidth of each legend object
            leg = ax.get_legend()
            for line in leg.get_lines():
                line.set_linewidth(4.0)

            ax.xaxis.tick_bottom()
            ax.yaxis.tick_left()
            plt.tight_layout()
            fig.savefig(fname, facecolor = 'w', bbox_inches='tight')
            return plt, ax

    doPlot (z, "./paper/Figure_3.png")
    paragraph = document.add_paragraph('')
    document.add_picture("./paper/Figure_3.png", width=Inches(6.0))
    paragraph = document.add_paragraph('Figure S' + str(cFigNumber) + ": Relation of feature stability with the number of selected features.")
    cFigNumber = cFigNumber + 1

    # each dataset
    for d in dList:
        z = df.query("Dataset == @d")
        z = z.groupby(["Feature Selection", "N"]).median(["Stability"])
        doPlot (z, "./results/Stability_" + d + ".png")
        # add to doc
        document.add_page_break()
        paragraph = document.add_paragraph(d)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture("./results/Stability_"+d+".png", width=Inches(6.0))
        paragraph = document.add_paragraph('Figure S' + str(cFigNumber) + ": Relation of feature stability with the number of selected features on dataset " + str(d) + ".")
        #paragraph = document.add_paragraph(" ")
        cFigNumber = cFigNumber + 1

    document.add_page_break()
    document.save('./paper/Supplemental_2.docx')
    plt.close('all')
    pass



def getSimilarity (cur, best):
    u = getFPattern (best)
    v = getFPattern (cur)
    SC = stability (u, v)
    return SC



# cross similairty of a pattern
def correlation (X, fu, fv):
    assert (len(fu) == len(fv))
    cu = []
    idxU = np.where(fu == 1.0)[0]
    idxV = np.where(fv == 1.0)[0]
    for u in idxU:
        cv = []
        for v in idxV:
            corr, pval = pearsonr (X.values[:,u], X.values[:,v])
            cv.append(np.abs(corr))
        cu.append(np.max(cv))
    CS = np.mean(cu)
    return CS



# called ucorr in supplemental
def getRawCorrelation (cur, best, X):
    assert (best["Dataset"] == cur["Dataset"])

    patBest = getFPattern (best)
    patCur = getFPattern (cur)
    assert (len(patBest) == len(patCur))
    m = len(patBest)
    CS = 0
    for i in range(m):
        for j in range(m):
            coef = correlation (X, patBest[i,:], patCur[j,:])
            CS = CS + coef
    CS = 1/(m*m)*CS
    return CS


def getCorrelation (cur, best, X):
    assert (best["Dataset"] == cur["Dataset"])
    CSp = getRawCorrelation (cur, best, X)
    CSm = getRawCorrelation (best, cur, X)
    return (CSp + CSm)/2.0



def highlight_cell(x,y, ax=None, **kwargs):
    rect = plt.Rectangle((x-.5, y-.5), 1,1, fill=False, **kwargs)
    ax = ax or plt.gca()
    ax.add_patch(rect)
    return rect



def plot_Tables (dList, results, cType = None):
    # set params
    if cType == "Stability":
        sFile = "Supplemental 3"
    if cType == "Similarity":
        sFile = "Supplemental 4"
    if cType == "Correlation":
        sFile = "Supplemental 5"
    # create supplemental
    global document;
    global cFigNumber;
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading(sFile)
    document.add_paragraph(' ')

    document.add_heading("Feature " + cType, level = 2)
    document.add_paragraph(' ')

    rMatList = getDeLongTest (dList, results)
    sscMatList = getSSC (dList, results)

    # prepare data
    for d in dList:
        # ssc contains only stat sim models
        z = sscMatList[d]
        rMat = rMatList[d]
        rMat = rMat.round(3)
        scMat = rMat.copy()
        strMat = rMat.copy()
        strMat = strMat.astype( dtype = "str")

        # get best one
        aTable = results.query("Dataset == @d")
        aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
        best = aTable.iloc[0].copy()
        bestc = best["Clf"]
        bestf = best["FSel"]

        for (i, (idx, row)) in enumerate(z.iterrows()):
            c = row["Classifier"]
            f = row["Feature Selection"]
            nF = row["nFeatures"]
            auc = row["AUC"]
            scMat.at[c,f] = row[cType]
            k = str(np.round(scMat.at[c, f], 2))
            if k == "-0.0":
                k = "0.0"
            strMat.at[c,f] =  r'\huge{' + k + "}\n\Large{AUC:" + str(auc) + " (p=" + str(rMat.at[c,f]) + ")"+ "\n" + "\large{\#Features: " + str(nF) + "}"
            strMat.at[c,f]

        if 1 == 1:
            plt.rc('text', usetex=True)
            plt.rcParams.update({
                "font.family": "sans-serif",
                "font.sans-serif": ["Arial"]})
            plt.rcParams['text.usetex'] = True
            plt.rcParams['text.latex.preamble'] = r'''
                \usepackage{mathtools}
                \usepackage{helvet}
                \renewcommand{\familydefault}{\sfdefault}        '''

            DPI  = 300
            fig, ax = plt.subplots(figsize = (17,14), dpi = DPI)
            sns.set(style='white')
            #ax = sns.heatmap(scMat, annot = cMat, cmap = "Blues", fmt = '', annot_kws={"fontsize":21}, linewidth = 2.0, linecolor = "black")
            dx = np.asarray(scMat, dtype = np.float64)
            pal = sns.light_palette("#8888bb", reverse=False, as_cmap=True)
            tnorm = colors.TwoSlopeNorm(vmin=0.0, vcenter=0.5, vmax=1.0)
            ax.imshow(dx, cmap=pal, norm = tnorm, interpolation='nearest', aspect = 0.49)

            # Major ticks
            mh, mw = scMat.shape
            ax.set_xticks(np.arange(0, mw, 1))
            ax.set_yticks(np.arange(0, mh, 1))

            # Minor ticks
            ax.set_xticks(np.arange(-.5, mw, 1), minor=True)
            ax.set_yticks(np.arange(-.5, mh, 1), minor=True)

            # Gridlines based on minor ticks
            ax.grid(which='minor', color='black', linestyle='-', linewidth=2)

            for i, c in enumerate(scMat.index):
                for j, f in enumerate(scMat.keys()):
                    if rMat.at[c,f] < 0.05:
                        ax.text(j, i, '  ', ha="center", va="center", color="k", fontsize = 12)
                    elif scMat.at[c,f] < -0.95:
                        ax.text(j, i, '\huge{N/A}', ha="center", va="center", color="k", fontsize = 12)
                    else:
                        ax.text(j, i, strMat.at[c, f],    ha="center", va="center", color="k", fontsize = 12)
            plt.tight_layout()

            bestcidx = list(scMat.keys()).index(bestf)
            bestfidx = list(scMat.index).index(bestc)
            highlight_cell(bestcidx, bestfidx, color="royalblue", linewidth=10)

            ax.set_xticklabels(rMat.keys(), rotation = 45, ha = "right", fontsize = 22)
            ax.set_yticklabels(strTrans(rMat.index), rotation = 0, ha = "right", fontsize = 22)
            ax.yaxis.set_tick_params ( labelsize= 22)
            fig.savefig("./results/Table_" + cType + "_"+d+".png", facecolor = 'w', bbox_inches='tight')
            paragraph = document.add_paragraph(cType + " on " + d)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_picture("./results/Table_" + cType + "_" + d+".png", width=Inches(6.0))
            paragraph = document.add_paragraph('Figure S' + str(cFigNumber) + ": " + cType + " of the models on dataset " + str(d) + ". The best model is framed with a blue border, models that " + \
                "were significantly different to the best model are not shown. Statistical significance was tested using a DeLong test.")
            cFigNumber = cFigNumber + 1
            document.add_page_break()

        plt.close('all')
    plt.rc('text', usetex=False)

    document.save("./paper/" + sFile.replace(" ", "_") + ".docx")
    plt.close('all')
    pass



def getSSC (dList, results):
    rMatList = getDeLongTest (dList, results)
    if os.path.exists("./results/ssc.feather") == False:
        sscMatList = {}
        for fidx, d in enumerate(dList):
            X, y = datasets[d]

            rMat = rMatList[d]
            rMat = rMat.round(3)
            fSels = sorted(list(set(results["FSel"].values)))
            nList = sorted(list(set(results["nFeatures"].values)))
            clfs = sorted(list(set(results["Clf"].values)))

            # get best overall
            aTable = results.query("Dataset == @d")
            aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
            best = aTable.iloc[0].copy()

            fTable = []
            for c in clfs:
                for f in fSels:
                    if rMat.at[c,f] > 0.05:
                        aTable = results.query("Dataset == @d and FSel == @f and Clf == @c")
                        aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
                        cur = aTable.iloc[0].copy()

                        sim = getSimilarity (cur, best)
                        stab = getStability (cur)
                        corr = getCorrelation (cur, best, X)
                        fTable.append({"Feature Selection": f, "Classifier": c, "AUC": round(cur["AUC"],2),
                                "nFeatures": cur["nFeatures"],
                                "Stability": round(stab,2),
                                "Similarity": round(sim, 2),
                                "Correlation": round(corr,2)})
            sscMatList[d] = pd.DataFrame(fTable)
        print ("Pickling SSC results")
        pickle.dump (sscMatList, open("./results/ssc.feather","wb"))
    else:
        print ("Restoring SSC results")
        sscMatList = pickle.load(open("./results/ssc.feather", "rb"))
    return sscMatList




def table_ZZ (dList, results):
    rMatList = getDeLongTest (dList, results)
    sscMatList = getSSC (dList, results)

    # prepare data
    df = []
    for d in dList:
        # ssc contains only stat sim models
        z = sscMatList[d]
        # did not save the best model, find it
        aTable = results.query("Dataset == @d")
        aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
        best = aTable.iloc[0].copy()
        for (i, (idx, row)) in enumerate(z.iterrows()):
            if row["Feature Selection"] == best["FSel"] and row["Classifier"] == best["Clf"] and row["nFeatures"] == best["nFeatures"]:
                continue
            df.append({"Dataset": d, "Stability": row["Stability"],  "Similarity": row["Similarity"], "Correlation": row["Correlation"] })

    df = pd.DataFrame(df)
    rMat = df.groupby(["Dataset"]).mean()
    minMat = df.groupby(["Dataset"]).min()
    maxMat = df.groupby(["Dataset"]).max()
    minMat = minMat.round(2).astype(str)
    maxMat = maxMat.round(2).astype(str)

    # labels need range
    labels = rMat.copy()
    labels = labels.round(2).astype(str)

    for c in list(labels.index):
        for f in list(labels.keys()):
            labels.at[c,f] = labels.at[c,f] + "\n(" + minMat.at[c,f] + "-" + maxMat.at[c,f] + ")"
    labels = np.array(labels)


    DPI = 300
    if 1 == 1:
        fig, ax = plt.subplots(figsize = (10,15), dpi = DPI)
        sns.set(style='white')
        sns.heatmap(rMat,  annot = labels, cmap = "Reds", fmt = '', annot_kws={"fontsize":21}, cbar = False)
        ax.set_xticklabels(rMat.keys(), rotation = 45, ha = "right", fontsize = 21)
        ax.set_yticklabels(rMat.index, rotation = 0, ha = "right", fontsize = 21)
        ax.yaxis.set_tick_params ( labelsize= 21)
        ax.set_xlabel ("", fontsize = 19)
        ax.set_ylabel ("", fontsize = 19)
        ax.set_title("", fontsize = 24)
        plt.tight_layout()
        fig.savefig("./paper/Figure_4.png", facecolor = 'w')



def plot_ZZ (dList, results):
    rMatList = getDeLongTest (dList, results)
    sscMatList = getSSC (dList, results)

    def doPlot (z, v, fname, ax):
        sns.set(style='white')

        # prepare data
        df = []
        z = z.sort_values(["Stability"], ascending = False).reset_index(drop = True)
        for (i, (idx, row)) in enumerate(z.iterrows()):
            p = v.at[row["Classifier"], row["Feature Selection"]]
            if p >= 0.05:
                df.append({"Value": row["Stability"], "Type": "Stability", "Index": idx})
                df.append({"Value": row["Similarity"], "Type": "Similarity", "Index": idx})
                df.append({"Value": row["Correlation"], "Type": "Correlation", "Index": idx})
        df = pd.DataFrame(df)
        #z["Value"] = z["Value"].replace(-1, np.inf)
        palette =         sns.color_palette("hls", 17)[4::6]
        try:
            line=sns.lineplot(x="Index", y="Value",hue="Type", palette = palette, marker="o", data=df, linewidth = 4, ax = ax, legend =None)
        except Exception as e:
            print(z)
            print(z.head())
            raise(e)
        #ax = ax[0][0]
        ax.xaxis.set_tick_params(labelsize=20)
        ax.yaxis.set_tick_params(labelsize=20)
        ax.set_ylabel('Correlation', fontsize = 22, labelpad = 12)
        ax.set_xlabel('Model', fontsize= 22, labelpad = 12)
        ax.set_title(str(d), fontsize="24", fontweight="bold")
        ax.set_ylim(0, 1)

        ax.xaxis.tick_bottom()
        ax.yaxis.tick_left()
        extent = full_extent(ax).transformed(fig.dpi_scale_trans.inverted())
        # extent = ax.get_tightbbox(fig.canvas.renderer).transformed(fig.dpi_scale_trans.inverted())
        fig.savefig(fname, facecolor = 'w', bbox_inches=extent)
        return plt, ax

    DPI  = 200
    fig, ax = plt.subplots(4,4, figsize = (20, 20), dpi = DPI)
    for fidx, d in enumerate(dList):
        fidx_y= fidx % 4
        fidx_x = fidx//4
        z = sscMatList[d]
        v = rMatList[d] # fixme: add p value to sscmatlist....
        doPlot (z, v, "./results/Overview_" + d + ".png", ax = ax[fidx_x][fidx_y])

    # add legend to bottom right
    for x, y in [(3,3), (3,2)]:
        ax[x][y] .spines['right'].set_visible(False)
        ax[x][y] .spines['top'].set_visible(False)
        ax[x][y] .spines['bottom'].set_visible(False)
        ax[x][y] .spines['left'].set_visible(False)
        ax[x][y] .spines['left'].set_visible(False)
        ax[x][y].xaxis.set_visible(False)
        ax[x][y].yaxis.set_visible(False)

    palette =         sns.color_palette("hls", 17)[4::6]
    labels = ["Stability", "Similarity", "Correlation"]
    handles = [matplotlib.patches.Patch(color=x, label=labels[v]) for v,x in enumerate(palette)]
    # Create legend
    ax[3][3].legend(handles=handles, loc = "lower right")
    # Get current axes object and turn off axis
    ax[3][3].set_axis_off()
    #ax[3][3].legend(loc = "lower right")
    plt.setp(ax[3][3].get_legend().get_texts(), fontsize='24') # for legend text
    plt.setp(ax[3][3].get_legend().get_title(), fontsize='24') # for legend title

    plt.tight_layout(pad=3.0)
    fig.savefig("./results/Figure_ZZ.png", facecolor = 'w', bbox_inches='tight')
    plt.close('all')
    plt.rc('text', usetex=False)
    pass



# mean AUC of all stat. Sig. models vs mean correlation (pre/post/f-corr)
def plot_TradeOff (dList, results):
    print ("Plotting trade off")
    rMatList = getDeLongTest (dList, results)
    sscMatList = getSSC (dList, results)

    # prepare data
    print ("Preparing data")
    df = []
    for d in dList:
        # ssc contains only stat sim models
        z = sscMatList[d]
        # did not save the best model, find it
        aTable = results.query("Dataset == @d")
        aTable = aTable.sort_values("AUC", ascending = False).reset_index(drop = True).copy()
        best = aTable.iloc[0].copy()
        for (i, (idx, row)) in enumerate(z.iterrows()):
            if row["Feature Selection"] == best["FSel"] and row["Classifier"] == best["Clf"] and row["nFeatures"] == best["nFeatures"]:
                continue
            df.append({"Dataset": d, "AUC": row["AUC"], "Stability": row["Stability"],  "Similarity": row["Similarity"], "Correlation": row["Correlation"] })

    df = pd.DataFrame(df)
    count = df.groupby(["Dataset"]).count()["AUC"]
    df = df.groupby(["Dataset"]).mean()
    df["Count"] = count

    print ("Plotting")
    def doPlot(df, d, v1, v2):
        for ctype in [v2]:
            spfList = df[[v1, ctype]]
            R, pval = pearsonr(*zip (*spfList.values))
            R2 = R*R
            print (R, pval)

            # fSels = [z[0] for z in spfList.index]
            # dSets = [z[1] for z in spfList.index]

            x, y = zip(*spfList.values)
            p, cov = np.polyfit(x, y, 1, cov=True)                     # parameters and covariance from of the fit of 1-D polynom.
            y_model = equation(p, x)                                   # model using the fit parameters; NOTE: parameters here are coefficients

            # Statistics
            n =  len(x)                                          # number of observations
            ps = p.size                                                 # number of parameters
            dof = n - ps                                                # degrees of freedom
            t = stats.t.ppf(0.975, n - ps)                              # used for CI and PI bands

            # Estimates of Error in Data/Model
            resid = y - y_model
            chi2 = np.sum((resid / y_model)**2)                        # chi-squared; estimates error in data
            chi2_red = chi2 / dof                                      # reduced chi-squared; measures goodness of fit
            s_err = np.sqrt(np.sum(resid**2) / dof)                    # standard deviation of the error


            # plot
            if 1 == 1:
                DPI = 300
                fig, ax = plt.subplots(figsize = (10, 10), dpi = DPI)
                # sns.scatterplot (x = x,y = y,  ax = ax)
                sns.scatterplot (x = v1, y = ctype,  data=df,  ax = ax, s = 50, color = ".0")
                ax.plot(x, y_model, "-", color="0.1", linewidth=1.5, alpha=1.0, label="Fit")

                x2 = np.linspace(np.min(x), np.max(x), 100)
                y2 = equation(p, x2)

                # Confidence Interval (select one)
                plot_ci_manual(t, s_err, n, x, x2, y2, ax=ax)
                #plot_ci_bootstrap(x, y, resid, ax=ax)

                # Prediction Interval
                pi = t * s_err * np.sqrt(1 + 1/n + (x2 - np.mean(x))**2 / np.sum((x - np.mean(x))**2))
                ax.fill_between(x2, y2 + pi, y2 - pi, color="None", linestyle="--")
                # ax.plot(x2, y2 - pi, "--", color="0.5", label="95% Prediction Limits")
                # ax.plot(x2, y2 + pi, "--", color="0.5")


                # Figure Modifications --------------------------------------------------------
                # Borders
                ax.spines["top"].set_color("0.5")
                ax.spines["bottom"].set_color("0.5")
                ax.spines["left"].set_color("0.5")
                ax.spines["right"].set_color("0.5")
                ax.get_xaxis().set_tick_params(direction="out")
                ax.get_yaxis().set_tick_params(direction="out")
                ax.xaxis.tick_bottom()
                ax.yaxis.tick_left()
                #ax.invert_xaxis()

                # Labels
                #plt.title("Fit Plot for Weight", fontsize="14", fontweight="bold")

                plt.xticks(fontsize=20)
                plt.yticks(fontsize=20)
                if v2 == "Correlation":
                    plt.ylabel('Mean Correlation', fontsize = 22, labelpad = 12)
                if v2 == "Stability":
                    plt.ylabel('Mean Stability', fontsize = 22, labelpad = 12)
                if v2 == "Similarity":
                    plt.ylabel('Mean Similarity', fontsize = 22, labelpad = 12)
                if v2 == "Count":
                    plt.ylabel('Number of stat. similar Models', fontsize = 22, labelpad = 12)

                if v1 == "AUC":
                    plt.xlabel('Mean AUC-ROC', fontsize= 22, labelpad = 12)
                if v1 == "Correlation":
                    plt.xlabel('Mean Correlation', fontsize= 22, labelpad = 12)
                if v1 == "Sample Size":
                    plt.xlabel('Sample Size', fontsize = 22, labelpad = 12)
                    ax.set_xticks([50,250,500,750])

                right = 0.95
                ypos = 0.07 #0.93s
                legtext = ''
                if len(legtext ) > 0:
                    ypos = 0.07
                    legtext=legtext+"\n"

                plt.rcParams.update({
                    "text.usetex": True,
                    "font.family": "sans-serif",
                    "font.sans-serif": ["Helvetica"]})
                legpost = ''
                bbox_props = dict(fc="w", ec="0.5", alpha=0.9)
                pTxt = (' = {0:0.2f} ($p$ = {1:0.3f})').format(R2, pval)
                plt.text (right, ypos,
                          (legtext +  "$R^2$" + pTxt),
                          horizontalalignment='right',
                          size = 24, bbox  = bbox_props,
                          transform = ax.transAxes)
                plt.rcParams.update({
                    "text.usetex": False,
                    "font.family": "sans-serif",
                    "font.sans-serif": ["Helvetica"]})

                #ax.set_title("Stability vs AUC (" + d + ")", fontsize = 28)
                print ("Bias for", d)
                fig.tight_layout()
                #fig.savefig("./results/AUC_vs_" + ctype + "_"+ d + ".png", facecolor = 'w')
                if d == "all":
                    if v2 == "Correlation" and v1 =="AUC":
                        fig.savefig("./results/Figure_3A.png", facecolor = 'w')
                    if v2 == "Count" and v1 == "AUC":
                        fig.savefig("./results/Figure_3B.png", facecolor = 'w')
                    if v2 == "Stability" and v1 == "AUC":
                        fig.savefig("./results/Figure_3C.png", facecolor = 'w')
                    if v2 == "Similarity" and v1 == "AUC":
                        fig.savefig("./results/Figure_3D.png", facecolor = 'w')



    doPlot (df, d = "all", v1 = "AUC", v2 = "Correlation")
    doPlot (df, d = "all", v1 = "AUC", v2 = "Count")
    doPlot (df, d = "all", v1 = "AUC", v2 = "Stability")
    doPlot (df, d = "all", v1 = "AUC", v2 = "Similarity")

    plt.close('all')
    pass




def addText (finalImage, text = '', org = (0,0), fontFace = '', fontSize = 12, color = (255,255,255)):
     # Convert the image to RGB (OpenCV uses BGR)
     #tmpImg = cv2.cvtColor(finalImage, cv2.COLOR_BGR2RGB)
     tmpImg = finalImage
     pil_im = Image.fromarray(tmpImg)
     draw = ImageDraw.Draw(pil_im)
     font = ImageFont.truetype(fontFace + ".ttf", fontSize)
     draw.text(org, text, font=font, fill = color)
     #tmpImg = cv2.cvtColor(np.array(pil_im), cv2.COLOR_RGB2BGR)
     tmpImg = np.array(pil_im)
     return (tmpImg.copy())



def addBorder (img, pos, thickness):
    if pos == "H":
        img = np.hstack([255*np.ones(( img.shape[0],int(img.shape[1]*thickness), 3), dtype = np.uint8),img])
    if pos == "V":
        img = np.vstack([255*np.ones(( int(img.shape[0]*thickness), img.shape[1], 3), dtype = np.uint8),img])
    return img


# just read both figures and merge them
def join_plots():
    fontFace = "Arial"

    imA = cv2.imread("./results/Figure_3B.png")
    imA = addText (imA, "(a)", (50,50), fontFace, 128, color=(0,0,0))

    imB = cv2.imread("./results/Figure_3C.png")
    imB = addText (imB, "(b)", (50,50), fontFace, 112, color= (0,0,0))

    imC = cv2.imread("./results/Figure_3D.png")
    imC = addText (imC, "(c)", (50,50), fontFace, 112, color=(0,0,0))

    imD = cv2.imread("./results/Figure_3A.png")
    imD = addText (imD, "(d)", (50,50), fontFace, 112, color= (0,0,0))
    #Image.fromarray(imD[::4,::4,:])
    imB = addBorder (imB, "H", 0.075)
    imgU = np.hstack([imA, imB])

    imD = addBorder (imD, "H", 0.075)
    imgL = np.hstack([imC, imD])

    imgL = addBorder (imgL, "V", 0.075)
    img = np.vstack([imgU, imgL])
    #Image.fromarray(img[::6,::6,:])

    cv2.imwrite("./paper/Figure_5.png", img)



def stabilityStats (dList, results):
    print ("Feature stability stats:")
    fTable = load( "./results/stability.joblib")

    # group by dataset and take mean
    df = pd.DataFrame(fTable)
    z = df.query("N == 1")
    print ("Of ", z.shape[0], "models selecting one feature.")
    zh = z.query("Stability >= 0.5")
    print ("...only ", zh.shape[0], "have corr > 0.5..")
    print ("This is ", zh.shape[0]/z.shape[0], "% of  the models.")

    # per dataset:
    for d in dList:
        z = df.query("N == 1 and Dataset == @d")
        zh = z.query("Stability >= 0.5")
        print ("For", d, ":", zh.shape[0], " models.", zh["Feature Selection"].values)
    print ("\n\n")




def full_extent(ax, pad=0.0):
    """Get the full extent of an axes, including axes labels, tick labels, and
    titles."""
    # For text objects, we need to draw the figure first, otherwise the extents
    # are undefined.
    ax.figure.canvas.draw()
    items = ax.get_xticklabels() + ax.get_yticklabels()
    items += [ax, ax.title, ax.xaxis.label, ax.yaxis.label]
    items += [ax, ax.title]
    bbox = Bbox.union([item.get_window_extent() for item in items])
    return bbox.expanded(1.0 + pad, 1.0 + pad)



def doHistoPlot (z, d, fname, ax, color = None, range = False, fig = None):
    sns.set(style='white')
    line = sns.histplot(z, bins = 40,  color = color, ax = ax, stat = "probability")

    ax.xaxis.set_tick_params(labelsize=20)
    ax.yaxis.set_tick_params(labelsize=20)
    ax.set_ylabel('Probability', fontsize = 22, labelpad = 12)
    ax.set_xlabel('Correlation [Pearson]', fontsize= 22, labelpad = 12)
    ax.set_title(str(d), fontsize="24", fontweight="bold")

    if range == True:
        ax.set_xlim(-1, 1)

    ax.xaxis.tick_bottom()
    ax.yaxis.tick_left()
    extent = full_extent(ax).transformed(fig.dpi_scale_trans.inverted())
    # extent = ax.get_tightbbox(fig.canvas.renderer).transformed(fig.dpi_scale_trans.inverted())
    fig.savefig(fname, facecolor = 'w', bbox_inches=extent)
    return plt, ax




if __name__ == "__main__":
    print ("Hi.")

    # load datasets
    print ("Loading datasets")
    datasets = {}
    for fidx, d in enumerate(dList):
        data = eval (d+"().getData('./data/')")
        y = data["Target"]
        X = data.drop(["Target"], axis = 1)
        X, y = preprocessData (X, y)
        datasets[d] = (X, y)

    # obtain results
    print ("Generating results")
    results = getResults (dList)
    _ = getDeLongTest (dList, results)

    # plot datasets
    plot_DataHisto (dList)

    # first do counting
    countModels (dList, results)
    plot_modelAUCs (dList, results)

    getSSC (dList, results)
    plot_ZZ(dList, results)
    table_ZZ (dList, results)

    # Stability of each model
    plot_Stability_Curves (dList, results)
    stabilityStats (dList, results)

    # Similarity between the models
    plot_Tables (dList, results, cType = "Stability")
    plot_Tables (dList, results, cType = "Similarity")
    plot_Tables (dList, results, cType = "Correlation")

    # Trade-off
    plot_TradeOff (dList, results)
    join_plots()


#
